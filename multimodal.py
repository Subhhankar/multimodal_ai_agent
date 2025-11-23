import fitz  # PyMuPDF
from langchain_core.documents import Document
from transformers import CLIPProcessor, CLIPModel
from PIL import Image
import torch
import numpy as np
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from langchain_core.messages import HumanMessage
from sklearn.metrics.pairwise import cosine_similarity
import os
import base64
import io
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv
import pandas as pd
import docx
from pathlib import Path
from typing import List, Dict, Tuple, Union
import mimetypes
import pytesseract
from pytesseract import Output
import faiss
import json

load_dotenv()

class MultimodalDocumentProcessor:
    """
    Universal document processor supporting:
    - PDFs (with text and images)
    - Images (PNG, JPG, JPEG, GIF, BMP, TIFF)
    - Text files (TXT, MD)
    - Word documents (DOCX)
    - Excel files (XLSX, XLS, CSV)
    """
    
    def __init__(self, google_api_key: str = None, enable_ocr: bool = True, store_as_base64: bool = False, use_faiss_index: bool = True):
        # Initialize CLIP model
        self.clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
        self.clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
        self.clip_model.eval()
        
        # Text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500, 
            chunk_overlap=100
        )
        
        # Storage
        self.all_docs = []
        self.all_embeddings = []
        self.image_data_store = {}
        
        # OCR settings
        self.enable_ocr = enable_ocr
        self.store_as_base64 = store_as_base64
        
        # FAISS indexing
        self.use_faiss_index = use_faiss_index
        self.faiss_index = None
        self.embedding_dimension = 512  # CLIP embedding size
        
        # API key setup
        if google_api_key:
            os.environ["GOOGLE_API_KEY"] = google_api_key
        
        # Initialize Gemini model
        self.llm = None
        if os.environ.get("GOOGLE_API_KEY"):
            try:
                self.llm = ChatGoogleGenerativeAI(
                    model="gemini-2.0-flash",
                    temperature=0.3,
                    max_tokens=2048
                )
            except Exception as e:
                print(f"⚠️  Could not initialize Gemini: {e}")
                print("    Set GOOGLE_API_KEY to use Q&A features")
        
        # Supported file extensions
        self.supported_extensions = {
            'pdf': self._process_pdf,
            'png': self._process_image,
            'jpg': self._process_image,
            'jpeg': self._process_image,
            'gif': self._process_image,
            'bmp': self._process_image,
            'tiff': self._process_image,
            'txt': self._process_text,
            'md': self._process_text,
            'docx': self._process_docx,
            'xlsx': self._process_excel,
            'xls': self._process_excel,
            'csv': self._process_csv
        }
    
    def extract_text_from_image(self, image: Image.Image) -> str:
        """Extract text from image using OCR"""
        if not self.enable_ocr:
            return ""
        
        try:
            # Check if pytesseract is available
            text = pytesseract.image_to_string(image, config='--psm 6')
            return text.strip()
        except pytesseract.TesseractNotFoundError:
            # Tesseract not installed - disable OCR silently
            if self.enable_ocr:
                print("⚠️  Tesseract OCR not found. OCR disabled. To enable:")
                print("   - Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki")
                print("   - Mac: brew install tesseract")
                print("   - Linux: sudo apt-get install tesseract-ocr")
                self.enable_ocr = False  # Disable to avoid repeated warnings
            return ""
        except Exception as e:
            # Silent failure for other OCR errors
            return ""
    
    def image_to_base64(self, image: Image.Image) -> str:
        """Convert PIL Image to base64 string"""
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        img_bytes = buffered.getvalue()
        return base64.b64encode(img_bytes).decode('utf-8')
    
    def base64_to_image(self, base64_str: str) -> Image.Image:
        """Convert base64 string to PIL Image"""
        img_bytes = base64.b64decode(base64_str)
        return Image.open(io.BytesIO(img_bytes))
    
    def store_image(self, image: Image.Image, img_id: str):
        """Store image in the selected format"""
        if self.store_as_base64:
            self.image_data_store[img_id] = self.image_to_base64(image)
        else:
            self.image_data_store[img_id] = image
    
    def get_image(self, img_id: str) -> Image.Image:
        """Retrieve image from storage"""
        stored_data = self.image_data_store.get(img_id)
        if stored_data is None:
            return None
        
        if self.store_as_base64:
            return self.base64_to_image(stored_data)
        else:
            return stored_data
    
    def build_faiss_index(self):
        """Build or rebuild FAISS index from current embeddings"""
        if not self.all_embeddings:
            print("No embeddings to index")
            return
        
        embeddings_array = np.array(self.all_embeddings).astype('float32')
        
        # Create FAISS index
        # Using IndexFlatIP (Inner Product) for cosine similarity
        # For larger datasets, consider IndexIVFFlat or IndexHNSWFlat
        self.faiss_index = faiss.IndexFlatIP(self.embedding_dimension)
        
        # Normalize embeddings for cosine similarity
        faiss.normalize_L2(embeddings_array)
        
        # Add embeddings to index
        self.faiss_index.add(embeddings_array)
        
        print(f"✓ FAISS index built with {self.faiss_index.ntotal} vectors")
    
    def add_to_index(self, embedding: np.ndarray):
        """Add a single embedding to the FAISS index"""
        if not self.use_faiss_index:
            return
        
        if self.faiss_index is None:
            # Create new index
            self.faiss_index = faiss.IndexFlatIP(self.embedding_dimension)
        
        # Normalize and add
        embedding_normalized = embedding.astype('float32').reshape(1, -1)
        faiss.normalize_L2(embedding_normalized)
        self.faiss_index.add(embedding_normalized)
    
    def embed_image(self, image_data: Union[str, Image.Image]) -> np.ndarray:
        """Embed image using CLIP"""
        try:
            if isinstance(image_data, str):
                image = Image.open(image_data).convert("RGB")
            else:
                image = image_data.convert("RGB")
            
            # Validate image dimensions
            width, height = image.size
            if width < 10 or height < 10:
                # Image too small, create dummy embedding
                print(f"⚠️  Skipping tiny image ({width}x{height})")
                return np.zeros(self.embedding_dimension, dtype=np.float32)
            
            # Resize if image is too large (prevent memory issues)
            max_size = 2000
            if max(width, height) > max_size:
                ratio = max_size / max(width, height)
                new_size = (int(width * ratio), int(height * ratio))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            inputs = self.clip_processor(images=image, return_tensors="pt")
            with torch.no_grad():
                features = self.clip_model.get_image_features(**inputs)
                features = features / features.norm(dim=-1, keepdim=True)
                return features.squeeze().numpy()
                
        except Exception as e:
            print(f"⚠️  Image embedding error: {e}")
            # Return zero vector as fallback
            return np.zeros(self.embedding_dimension, dtype=np.float32)
    
    def embed_text(self, text: str) -> np.ndarray:
        """Embed text using CLIP"""
        inputs = self.clip_processor(
            text=text,
            return_tensors='pt',
            padding=True,
            truncation=True,
            max_length=77
        )
        with torch.no_grad():
            features = self.clip_model.get_text_features(**inputs)
            features = features / features.norm(dim=-1, keepdim=True)
            return features.squeeze().numpy()
    
    def _process_pdf(self, file_path: str) -> List[Tuple[Document, np.ndarray]]:
        """Process PDF files with text and images"""
        results = []
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text
            text = page.get_text()
            if text.strip():
                chunks = self.text_splitter.split_text(text)
                for i, chunk in enumerate(chunks):
                    doc_obj = Document(
                        page_content=chunk,
                        metadata={
                            'source': file_path,
                            'page': page_num,
                            'chunk': i,
                            'type': 'text'
                        }
                    )
                    embedding = self.embed_text(chunk)
                    results.append((doc_obj, embedding))
            
            # Extract images
            image_list = page.get_images(full=True)
            for img_idx, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Convert to PIL Image
                    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                    
                    # Skip if image is too small or corrupted
                    width, height = image.size
                    if width < 10 or height < 10:
                        continue
                    
                    # Extract text from image using OCR
                    ocr_text = self.extract_text_from_image(image)
                    
                    # Store image
                    img_id = f"{file_path}_page{page_num}_img{img_idx}"
                    self.store_image(image, img_id)
                    
                    # Create document for image embedding
                    image_content = f"Image from page {page_num}"
                    if ocr_text:
                        image_content += f"\nExtracted text: {ocr_text}"
                    
                    doc_obj = Document(
                        page_content=image_content,
                        metadata={
                            'source': file_path,
                            'page': page_num,
                            'image_id': img_id,
                            'type': 'image',
                            'ocr_text': ocr_text
                        }
                    )
                    embedding = self.embed_image(image)
                    results.append((doc_obj, embedding))
                    
                    # If OCR extracted text, create additional text-based document
                    if ocr_text and len(ocr_text) > 20:  # Only if meaningful text
                        text_doc = Document(
                            page_content=ocr_text,
                            metadata={
                                'source': file_path,
                                'page': page_num,
                                'image_id': img_id,
                                'type': 'image_text',
                                'parent_image': img_id
                            }
                        )
                        text_embedding = self.embed_text(ocr_text)
                        results.append((text_doc, text_embedding))
                
                except Exception as e:
                    print(f"⚠️  Error processing image {img_idx} on page {page_num}: {str(e)[:100]}")
                    continue
        
        doc.close()
        return results
    
    def _process_image(self, file_path: str) -> List[Tuple[Document, np.ndarray]]:
        """Process standalone image files"""
        image = Image.open(file_path).convert("RGB")
        
        # Extract text using OCR
        ocr_text = self.extract_text_from_image(image)
        
        img_id = file_path
        self.store_image(image, img_id)
        
        # Create document with OCR text
        image_content = f"Image file: {Path(file_path).name}"
        if ocr_text:
            image_content += f"\nExtracted text: {ocr_text}"
        
        doc_obj = Document(
            page_content=image_content,
            metadata={
                'source': file_path,
                'image_id': img_id,
                'type': 'image',
                'ocr_text': ocr_text
            }
        )
        embedding = self.embed_image(image)
        results = [(doc_obj, embedding)]
        
        # Create additional text document if OCR found text
        if ocr_text and len(ocr_text) > 20:
            text_doc = Document(
                page_content=ocr_text,
                metadata={
                    'source': file_path,
                    'image_id': img_id,
                    'type': 'image_text',
                    'parent_image': img_id
                }
            )
            text_embedding = self.embed_text(ocr_text)
            results.append((text_doc, text_embedding))
        
        return results
    
    def _process_text(self, file_path: str) -> List[Tuple[Document, np.ndarray]]:
        """Process plain text files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        chunks = self.text_splitter.split_text(text)
        results = []
        for i, chunk in enumerate(chunks):
            doc_obj = Document(
                page_content=chunk,
                metadata={
                    'source': file_path,
                    'chunk': i,
                    'type': 'text'
                }
            )
            embedding = self.embed_text(chunk)
            results.append((doc_obj, embedding))
        
        return results
    
    def _process_docx(self, file_path: str) -> List[Tuple[Document, np.ndarray]]:
        """Process Word documents"""
        doc = docx.Document(file_path)
        results = []
        
        # Extract text from paragraphs
        full_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        chunks = self.text_splitter.split_text(full_text)
        
        for i, chunk in enumerate(chunks):
            doc_obj = Document(
                page_content=chunk,
                metadata={
                    'source': file_path,
                    'chunk': i,
                    'type': 'text'
                }
            )
            embedding = self.embed_text(chunk)
            results.append((doc_obj, embedding))
        
        # Extract images from Word document
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    image = Image.open(io.BytesIO(image_data)).convert("RGB")
                    
                    # Extract text using OCR
                    ocr_text = self.extract_text_from_image(image)
                    
                    img_id = f"{file_path}_img{i}"
                    self.store_image(image, img_id)
                    
                    # Create image document
                    image_content = "Image from Word document"
                    if ocr_text:
                        image_content += f"\nExtracted text: {ocr_text}"
                    
                    doc_obj = Document(
                        page_content=image_content,
                        metadata={
                            'source': file_path,
                            'image_id': img_id,
                            'type': 'image',
                            'ocr_text': ocr_text
                        }
                    )
                    embedding = self.embed_image(image)
                    results.append((doc_obj, embedding))
                    
                    # Create text document if OCR found text
                    if ocr_text and len(ocr_text) > 20:
                        text_doc = Document(
                            page_content=ocr_text,
                            metadata={
                                'source': file_path,
                                'image_id': img_id,
                                'type': 'image_text',
                                'parent_image': img_id
                            }
                        )
                        text_embedding = self.embed_text(ocr_text)
                        results.append((text_doc, text_embedding))
                except Exception as e:
                    print(f"Error processing image in Word doc: {e}")
        
        return results
    
    def _process_excel(self, file_path: str) -> List[Tuple[Document, np.ndarray]]:
        """Process Excel files"""
        df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets
        results = []
        
        for sheet_name, sheet_df in df.items():
            # Convert dataframe to text representation
            text = f"Sheet: {sheet_name}\n"
            text += sheet_df.to_string(index=False)
            
            chunks = self.text_splitter.split_text(text)
            for i, chunk in enumerate(chunks):
                doc_obj = Document(
                    page_content=chunk,
                    metadata={
                        'source': file_path,
                        'sheet': sheet_name,
                        'chunk': i,
                        'type': 'excel'
                    }
                )
                embedding = self.embed_text(chunk)
                results.append((doc_obj, embedding))
        
        return results
    
    def _process_csv(self, file_path: str) -> List[Tuple[Document, np.ndarray]]:
        """Process CSV files"""
        df = pd.read_csv(file_path)
        text = df.to_string(index=False)
        
        chunks = self.text_splitter.split_text(text)
        results = []
        
        for i, chunk in enumerate(chunks):
            doc_obj = Document(
                page_content=chunk,
                metadata={
                    'source': file_path,
                    'chunk': i,
                    'type': 'csv'
                }
            )
            embedding = self.embed_text(chunk)
            results.append((doc_obj, embedding))
        
        return results
    
    def process_file(self, file_path: str) -> bool:
        """
        Process a single file of any supported type
        Returns True if successful, False otherwise
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            print(f"Error: File {file_path} does not exist")
            return False
        
        extension = file_path.suffix.lower().lstrip('.')
        
        if extension not in self.supported_extensions:
            print(f"Error: File type .{extension} is not supported")
            print(f"Supported types: {', '.join(self.supported_extensions.keys())}")
            return False
        
        try:
            print(f"Processing {file_path.name}...")
            processor_func = self.supported_extensions[extension]
            results = processor_func(str(file_path))
            
            for doc, embedding in results:
                self.all_docs.append(doc)
                self.all_embeddings.append(embedding)
                
                # Add to FAISS index incrementally
                if self.use_faiss_index:
                    self.add_to_index(embedding)
            
            print(f"✓ Successfully processed {file_path.name} ({len(results)} chunks/items)")
            return True
            
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return False
    
    def process_directory(self, directory_path: str, recursive: bool = False):
        """Process all supported files in a directory"""
        directory = Path(directory_path)
        
        if not directory.exists() or not directory.is_dir():
            print(f"Error: {directory_path} is not a valid directory")
            return
        
        pattern = '**/*' if recursive else '*'
        files = [f for f in directory.glob(pattern) if f.is_file()]
        
        processed = 0
        for file_path in files:
            if self.process_file(str(file_path)):
                processed += 1
        
        print(f"\nProcessed {processed}/{len(files)} files")
    
    def process_multiple_files(self, file_paths: List[str]):
        """Process multiple files"""
        for file_path in file_paths:
            self.process_file(file_path)
    
    def get_embeddings_array(self) -> np.ndarray:
        """Get all embeddings as numpy array"""
        return np.array(self.all_embeddings)
    
    def search(self, query: str, top_k: int = 5, use_index: bool = None) -> List[Dict]:
        """Search through all processed documents"""
        if not self.all_embeddings:
            print("No documents processed yet!")
            return []
        
        # Determine which search method to use
        if use_index is None:
            use_index = self.use_faiss_index
        
        query_embedding = self.embed_text(query)
        
        if use_index and self.faiss_index is not None:
            # FAISS-based search (fast)
            return self._search_with_faiss(query_embedding, top_k)
        else:
            # Linear search (slow but always works)
            return self._search_linear(query_embedding, top_k)
    
    def _search_with_faiss(self, query_embedding: np.ndarray, top_k: int) -> List[Dict]:
        """Fast search using FAISS index"""
        # Normalize query embedding
        query_normalized = query_embedding.astype('float32').reshape(1, -1)
        faiss.normalize_L2(query_normalized)
        
        # Search
        similarities, indices = self.faiss_index.search(query_normalized, top_k)
        
        results = []
        for idx, similarity in zip(indices[0], similarities[0]):
            if idx < len(self.all_docs):  # Safety check
                doc = self.all_docs[idx]
                results.append({
                    'document': doc,
                    'similarity': float(similarity),
                    'content': doc.page_content,
                    'metadata': doc.metadata,
                    'index': int(idx)
                })
        
        return results
    
    def _search_linear(self, query_embedding: np.ndarray, top_k: int) -> List[Dict]:
        """Linear search (fallback method)"""
        embeddings_array = self.get_embeddings_array()
        
        similarities = cosine_similarity(
            query_embedding.reshape(1, -1),
            embeddings_array
        )[0]
        
        top_indices = np.argsort(similarities)[-top_k:][::-1]
        
        results = []
        for idx in top_indices:
            doc = self.all_docs[idx]
            results.append({
                'document': doc,
                'similarity': float(similarities[idx]),
                'content': doc.page_content,
                'metadata': doc.metadata,
                'index': int(idx)
            })
        
        return results
    
    def get_statistics(self) -> Dict:
        """Get statistics about processed documents"""
        stats = {
            'total_documents': len(self.all_docs),
            'total_images': len(self.image_data_store),
            'document_types': {},
            'sources': set(),
            'index_info': {
                'using_faiss': self.use_faiss_index,
                'index_built': self.faiss_index is not None,
                'indexed_vectors': self.faiss_index.ntotal if self.faiss_index else 0
            }
        }
        
        for doc in self.all_docs:
            doc_type = doc.metadata.get('type', 'unknown')
            stats['document_types'][doc_type] = stats['document_types'].get(doc_type, 0) + 1
            stats['sources'].add(doc.metadata.get('source', 'unknown'))
        
        stats['sources'] = list(stats['sources'])
        return stats
    
    def save_index(self, index_path: str = "faiss_index.bin", metadata_path: str = "metadata.pkl"):
        """Save FAISS index and metadata to disk"""
        if self.faiss_index is None:
            print("No index to save. Build index first.")
            return
        
        # Save FAISS index
        faiss.write_index(self.faiss_index, index_path)
        
        # Save metadata (documents and image store)
        import pickle
        metadata = {
            'docs': self.all_docs,
            'images': self.image_data_store,
            'embeddings': self.all_embeddings
        }
        with open(metadata_path, 'wb') as f:
            pickle.dump(metadata, f)
        
        print(f"✓ Index saved to {index_path}")
        print(f"✓ Metadata saved to {metadata_path}")
    
    def load_index(self, index_path: str = "faiss_index.bin", metadata_path: str = "metadata.pkl"):
        """Load FAISS index and metadata from disk"""
        import pickle
        
        # Load FAISS index
        self.faiss_index = faiss.read_index(index_path)
        
        # Load metadata
        with open(metadata_path, 'rb') as f:
            metadata = pickle.load(f)
        
        self.all_docs = metadata['docs']
        self.image_data_store = metadata['images']
        self.all_embeddings = metadata['embeddings']
        
        print(f"✓ Loaded {self.faiss_index.ntotal} vectors from {index_path}")
        print(f"✓ Loaded {len(self.all_docs)} documents")
        print(f"✓ Loaded {len(self.image_data_store)} images")
    
    def answer_question(
        self, 
        question: str, 
        top_k: int = 5,
        include_images: bool = True,
        max_images: int = 3,
        stream: bool = False
    ) -> Dict:
        """
        Answer a question using retrieved context and Gemini 2.0 Flash
        
        Args:
            question: User's question
            top_k: Number of documents to retrieve
            include_images: Whether to send images to Gemini
            max_images: Maximum number of images to include
            stream: Whether to stream the response
        
        Returns:
            Dictionary with answer, sources, and metadata
        """
        if self.llm is None:
            return {
                'answer': 'Error: Gemini API not initialized. Set GOOGLE_API_KEY.',
                'sources': [],
                'error': True
            }
        
        # 1. Retrieve relevant documents
        results = self.search(question, top_k=top_k)
        
        if not results:
            return {
                'answer': 'No relevant documents found in the database.',
                'sources': [],
                'error': False
            }
        
        # 2. Prepare context and images
        text_context = []
        images_to_send = []
        sources = []
        
        for idx, result in enumerate(results):
            doc = result['document']
            similarity = result['similarity']
            
            # Collect text context
            if doc.metadata.get('type') in ['text', 'image_text', 'csv', 'excel']:
                text_context.append(f"[Source {idx+1}] {doc.page_content}")
                sources.append({
                    'index': idx + 1,
                    'source': doc.metadata.get('source', 'Unknown'),
                    'type': doc.metadata.get('type'),
                    'page': doc.metadata.get('page'),
                    'similarity': float(similarity)
                })
            
            # Collect images
            if include_images and doc.metadata.get('type') == 'image':
                image_id = doc.metadata.get('image_id')
                if image_id and len(images_to_send) < max_images:
                    image = self.get_image(image_id)
                    if image:
                        images_to_send.append({
                            'image': image,
                            'source': doc.metadata.get('source', 'Unknown'),
                            'page': doc.metadata.get('page'),
                            'ocr_text': doc.metadata.get('ocr_text', '')
                        })
                        sources.append({
                            'index': idx + 1,
                            'source': doc.metadata.get('source', 'Unknown'),
                            'type': 'image',
                            'page': doc.metadata.get('page'),
                            'similarity': float(similarity)
                        })
        
        # 3. Build prompt
        context_text = "\n\n".join(text_context) if text_context else "No text context available."
        
        prompt = f"""You are a helpful AI assistant answering questions based on provided documents and images.

CONTEXT FROM DOCUMENTS:
{context_text}

USER QUESTION:
{question}

INSTRUCTIONS:
- Answer the question using ONLY the information from the provided context and images
- Be specific and cite which source ([Source N]) you're using
- If the context doesn't contain enough information, say so clearly
- If images are provided, analyze them carefully and incorporate relevant details
- Be concise but comprehensive

ANSWER:"""
        
        # 4. Prepare message with text and images
        message_content = [{"type": "text", "text": prompt}]
        
        # Add images if available
        for img_data in images_to_send:
            try:
                # Convert PIL image to base64
                buffered = io.BytesIO()
                img_data['image'].save(buffered, format="PNG")
                img_base64 = base64.b64encode(buffered.getvalue()).decode()
                
                # Add image context
                img_context = f"\n\n[Image from {img_data['source']}"
                if img_data.get('page'):
                    img_context += f", page {img_data['page']}"
                if img_data.get('ocr_text'):
                    img_context += f" - OCR text: {img_data['ocr_text']}"
                img_context += "]"
                
                message_content.append({
                    "type": "image_url",
                    "image_url": f"data:image/png;base64,{img_base64}"
                })
                message_content.append({
                    "type": "text",
                    "text": img_context
                })
            except Exception as e:
                print(f"⚠️  Error processing image: {e}")
        
        # 5. Get response from Gemini
        try:
            response = self.llm.invoke([HumanMessage(content=message_content)])
            answer = response.content
            
            return {
                'answer': answer,
                'sources': sources,
                'num_sources': len(sources),
                'num_images': len(images_to_send),
                'error': False
            }
        
        except Exception as e:
            return {
                'answer': f'Error getting response from Gemini: {str(e)}',
                'sources': sources,
                'error': True
            }
    
    def chat(self, question: str, top_k: int = 5, include_images: bool = True):
        """
        Simplified chat interface for asking questions
        
        Usage:
            processor.chat("What is the revenue for Q3?")
        """
        result = self.answer_question(question, top_k=top_k, include_images=include_images)
        
        print("\n" + "="*80)
        print(f"QUESTION: {question}")
        print("="*80)
        
        if result.get('error'):
            print(f"\n❌ {result['answer']}\n")
        else:
            print(f"\n📝 ANSWER:\n{result['answer']}\n")
            
            if result.get('sources'):
                print(f"\n📚 SOURCES ({result['num_sources']} documents")
                if result.get('num_images', 0) > 0:
                    print(f"           {result['num_images']} images):")
                else:
                    print("):")
                    
                for source in result['sources']:
                    source_str = f"  [{source['index']}] {source['source']}"
                    if source.get('page') is not None:
                        source_str += f" (page {source['page']})"
                    source_str += f" - {source['type']} - similarity: {source['similarity']:.3f}"
                    print(source_str)
        
        print("="*80 + "\n")
        
        return result
